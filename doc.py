import json
from typing import Any, Dict, List, Optional
from datetime import datetime, timezone

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


class DocReportGenerator:
    def __init__(self, api_url: str, output_path: str = "report.docx", headers: Optional[Dict[str, str]] = None, timeout: int = 15) -> None:
        self.api_url = api_url
        self.output_path = output_path
        self.headers = headers or {"Accept": "application/json"}
        self.timeout = timeout
        self.last_fetch_at: Optional[datetime] = None
        self.raw_data: Optional[Any] = None

    def fetch_data(self, params: Optional[Dict[str, Any]] = None) -> Any:
        try:
            resp = requests.get(self.api_url, headers=self.headers, params=params, timeout=self.timeout)
            resp.raise_for_status()
            try:
                data = resp.json()
            except json.JSONDecodeError as e:
                raise RuntimeError(f"Invalid JSON from API: {e}") from e
            self.raw_data = data
            self.last_fetch_at = datetime.now(timezone.utc)
            return data
        except requests.RequestException as e:
            raise RuntimeError(f"API request failed: {e}") from e

    def _normalize(self, data: Any) -> Dict[str, Any]:
        items: List[Dict[str, Any]] = []
        meta: Dict[str, Any] = {}
        if isinstance(data, list):
            items = data
            meta = {"source": self.api_url, "count": len(items)}
        elif isinstance(data, dict):
            items = [data]
            meta = {"source": self.api_url, "count": 1}
        else:
            items = []
            meta = {"source": self.api_url, "count": 0}
        normalized_items = []
        for i, item in enumerate(items, start=1):
            if not isinstance(item, dict):
                continue
            normalized_items.append({
                "index": i,
                "id": item.get("id", f"item-{i}"),
                "title": item.get("title") or "(untitled)",
                "status": "published",
                "created_at": "",
                "summary": item.get("body", ""),
            })
        return {"meta": meta, "items": normalized_items}

    def format_doc(self, normalized: Dict[str, Any]) -> Document:
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        title = doc.add_heading("API Report Summary", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
        meta = normalized.get("meta", {})
        p = doc.add_paragraph()
        p.add_run(f"Source: {meta.get('source', self.api_url)}\n").bold = True
        p.add_run(f"Items: {meta.get('count', 0)}\n")
        p.add_run(f"Generated: {generated_at}\n")
        if self.last_fetch_at:
            p.add_run(f"Fetched: {self.last_fetch_at.strftime('%Y-%m-%d %H:%M:%S UTC')}\n")
        items = normalized.get("items", [])
        if items:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid"
            hdr_cells = table.rows[0].cells
            headers = ["#", "ID", "Title", "Status", "Created"]
            for cell, text in zip(hdr_cells, headers):
                cell.text = text
            for row in items:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row["index"])
                row_cells[1].text = str(row["id"])
                row_cells[2].text = row["title"]
                row_cells[3].text = row["status"]
                row_cells[4].text = row["created_at"]
            doc.add_paragraph()
        for row in items:
            doc.add_heading(f"{row['index']}. {row['title']}", level=2)
            para = doc.add_paragraph()
            para.add_run(f"ID: {row['id']} | Status: {row['status']} | Created: {row['created_at']}\n").bold = True
            body = row.get("summary", "").strip() or "(no summary available)"
            run = doc.add_paragraph(body).runs[0]
            run.font.size = Pt(11)
        return doc

    def save(self, doc: Document) -> str:
        try:
            doc.save(self.output_path)
            return self.output_path
        except Exception as e:
            raise RuntimeError(f"Failed to save DOCX: {e}") from e

    def run(self, params: Optional[Dict[str, Any]] = None) -> str:
        data = self.fetch_data(params=params)
        normalized = self._normalize(data)
        doc = self.format_doc(normalized)
        return self.save(doc)


if __name__ == "__main__":
    api_url = "https://jsonplaceholder.typicode.com/posts"
    generator = DocReportGenerator(api_url=api_url, output_path="api_report.docx")
    try:
        path = generator.run()
        print(f"Report generated: {path}")
    except RuntimeError as err:
        print(f"Error: {err}")
